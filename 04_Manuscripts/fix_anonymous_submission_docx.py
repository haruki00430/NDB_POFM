#!/usr/bin/env python3
"""Post-process anonymous Health Policy submission DOCX for double-blind review."""

from __future__ import annotations

from pathlib import Path

from docx import Document

INPUT = Path("10Manuscript_health_policy_submission_v1_anonymous.docx")
OUTPUT = INPUT

AUTHOR_1_LINE = (
    "[Author 1]: Conceptualization, Data curation, Formal analysis, Investigation, "
    "Methodology, Software, Visualization, Writing – original draft, "
    "Writing – review and editing."
)
AUTHOR_2_LINE = (
    "[Author 2]: Conceptualization, Supervision, Writing – review and editing."
)

UK_TO_US_SPELLING = {
    "characterised": "characterized",
    "characterise": "characterize",
    "characterising": "characterizing",
}

# GitHub URL with username is omitted during double-blind review (Zenodo DOI only).
DATA_AVAIL_BLINDED = (
    "The NDB Open Data used in this analysis are publicly available from the "
    "Ministry of Health, Labour and Welfare of Japan "
    "(https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html). "
    "Analysis code is archived on Zenodo "
    "(https://doi.org/10.5281/zenodo.20756258). "
    "A public GitHub repository link will be provided upon acceptance."
)

DATA_AVAIL_IDENTIFYING_MARKERS = (
    "github.com/haruki00430",
    "Haruki Saito:",
    "Tetsuya Ohira:",
)


def _replace_paragraph_text(par, new_text: str) -> None:
    for run in par.runs:
        run.text = ""
    if par.runs:
        par.runs[0].text = new_text
    else:
        par.add_run(new_text)


def main() -> None:
    doc = Document(str(INPUT))
    changed: list[str] = []

    for par in doc.paragraphs:
        text = par.text.strip()

        if text.startswith("Objective:"):
            for run in par.runs:
                if run.text.startswith("Objective:"):
                    run.text = run.text.replace("Objective:", "Objectives:", 1)
                    changed.append("Abstract: Objective → Objectives")
                    break

        if text.startswith("Haruki Saito:"):
            _replace_paragraph_text(par, AUTHOR_1_LINE)
            changed.append("Author Contributions: [Author 1]")
        elif text.startswith("Tetsuya Ohira:"):
            _replace_paragraph_text(par, AUTHOR_2_LINE)
            changed.append("Author Contributions: [Author 2]")

        for run in par.runs:
            if run.text:
                updated = run.text
                for uk, us in UK_TO_US_SPELLING.items():
                    updated = updated.replace(uk, us)
                if updated != run.text:
                    run.text = updated
                    changed.append("UK → US spelling (characterize)")

        if any(m in text for m in DATA_AVAIL_IDENTIFYING_MARKERS) or (
            "zenodo" in text.lower() and "github.com" in text.lower()
        ):
            _replace_paragraph_text(par, DATA_AVAIL_BLINDED)
            changed.append("Data Availability blinded (Zenodo only)")

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    for item in dict.fromkeys(changed):
        print(f"  - {item}")


if __name__ == "__main__":
    main()
