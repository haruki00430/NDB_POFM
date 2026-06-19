#!/usr/bin/env python3
"""Create separate Title Page DOCX for Health Policy double-anonymized submission."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = '10b_TitlePage_health_policy.docx'

doc = Document()

def add(text, bold=False, size=12, align='left', space_after=0):
    p = doc.add_paragraph()
    p.alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

add('TITLE PAGE', bold=True, size=14, align='center', space_after=12)

add('Title:', bold=True, space_after=2)
add(('When Universal Coverage Meets Local Capacity: '
     'A Prefectural Ecological Analysis of Perioperative Oral Care in Japan'),
    size=12, space_after=12)

add('Running head:', bold=True, space_after=2)
add('Supply-Driven Implementation of Perioperative Oral Care in Japan', space_after=12)

add('Authors:', bold=True, space_after=2)
add('Haruki Saito¹ (Corresponding author)  |  ORCID: 0009-0009-7890-6068', space_after=2)
add('Tetsuya Ohira¹²  |  ORCID: 0000-0003-4532-7165', space_after=12)

add('Affiliations:', bold=True, space_after=2)
add(('¹ Department of Epidemiology, Fukushima Medical University School of Medicine, '
     'Fukushima, Japan'), space_after=2)
add(('² Radiation Medical Science Center for the Fukushima Health Management Survey, '
     'Fukushima Medical University, Fukushima, Japan'), space_after=12)

add('Corresponding author:', bold=True, space_after=2)
add('Haruki Saito, MD', space_after=2)
add(('Department of Epidemiology, Fukushima Medical University School of Medicine, '
     '1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan'), space_after=2)
add('Email: m211039@fmu.ac.jp  |  ORCID: 0009-0009-7890-6068', space_after=12)

add('Word count (main text: Introduction–Discussion):', bold=True, space_after=2)
add('Approximately 2,325 words (within the 4,000-word limit)', space_after=12)

add('Figures and Tables:', bold=True, space_after=2)
add('Main body: 2 tables (Table 1, Table 2) + 2 figures (Figure 3, Figure 6) = 4 total', space_after=2)
add('Appendix: Appendix Table A1; Appendix Figures A1, A2, A4, A5; Supplementary Figure S1', space_after=12)

add('Conflict of interest:', bold=True, space_after=2)
add('The authors declare no conflicts of interest.', space_after=12)

add('Funding:', bold=True, space_after=2)
add('None declared. No external funding was received for this study.', space_after=12)

add('Data availability:', bold=True, space_after=2)
add(('The NDB Open Data used in this analysis are publicly available from the Ministry of Health, '
     'Labour and Welfare of Japan.'), space_after=12)

add('Author Contributions (CRediT Taxonomy):', bold=True, space_after=2)
add(('Haruki Saito: Conceptualization, Data curation, Formal analysis, Investigation, Methodology, '
     'Software, Visualization, Writing – original draft, Writing – review and editing.'),
    space_after=2)
add('Tetsuya Ohira: Conceptualization, Supervision, Writing – review and editing.', space_after=12)

add('Ethics Statement:', bold=True, space_after=2)
add(('This study used publicly available aggregate data. Individual informed consent was not required, '
     'and institutional ethics review was not applicable in accordance with Japanese ethical guidelines '
     'for epidemiological research.'), space_after=12)

add('Data Availability Statement:', bold=True, space_after=2)
add(('The NDB Open Data used in this analysis are publicly available from the Ministry of Health, '
     'Labour and Welfare of Japan (https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html). '
     'Analysis code is openly available on GitHub (https://github.com/haruki00430/NDB_POFM) '
     'and archived on Zenodo (https://doi.org/10.5281/zenodo.20755988).'), space_after=12)

add('AI Use Disclosure:', bold=True, space_after=2)
add(('During the preparation and writing of this work, the authors used AI-assisted tools to support '
     'manuscript drafting and statistical analysis scripting. Cursor (Anysphere) and Google Antigravity '
     '(Google) were used for AI-assisted writing and Python code development. Large language models '
     'included Claude Sonnet 4.6 and Claude Opus 4.8 (Anthropic), GPT-5.5 (OpenAI), and Gemini 3 Pro '
     '(Google). These tools were used only for text drafting and code generation; no generative AI tools '
     'were used to create, alter, or process figures, images, or artwork. The authors reviewed and edited '
     'all AI-assisted outputs and take full responsibility for the integrity and accuracy of the final '
     'content.'), space_after=12)

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
