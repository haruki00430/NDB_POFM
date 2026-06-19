#!/usr/bin/env python3
"""Renumber references 20-22 to match Vancouver order of first citation in body."""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

MS_DIR = Path(__file__).resolve().parent
INPUT = MS_DIR / "10Manuscript_health_policy_submission_v1_anonymous.docx"
OUTPUT = INPUT

# Old list: 20=LISA 1995, 21=Anselin 1988, 22=LeSage 2009
# In-text first appearance: [21,22] then [20] → new: 20=1988, 21=LeSage, 22=LISA


def _backup(path: Path) -> Path:
    """Create a timestamped backup before modifying the manuscript."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = path.with_name(f"{path.stem}_backup_{timestamp}{path.suffix}")
    shutil.copy2(path, backup)
    return backup


def _replace_runs_text(par: Paragraph, new_text: str) -> None:
    for run in par.runs:
        run.text = ""
    if par.runs:
        par.runs[0].text = new_text
    else:
        par.add_run(new_text)


def _fix_in_text_citations(text: str) -> str:
    """Swap citation numbers 20/21/22 without disturbing other references."""
    text = text.replace("[21,22]", "<<CIT_20_21>>")
    text = text.replace("[22]", "<<CIT_21>>")
    text = text.replace("[20]", "[22]")
    text = text.replace("<<CIT_21>>", "[21]")
    return text.replace("<<CIT_20_21>>", "[20,21]")


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(f"Missing input: {INPUT}")

    backup = _backup(INPUT)
    print(f"Backup: {backup.name}")

    doc = Document(str(INPUT))
    in_refs = False
    ref_paras: dict[int, Paragraph] = {}

    for par in doc.paragraphs:
        t = par.text.strip()
        if t == "References":
            in_refs = True
            continue
        if in_refs:
            m = re.match(r"^(\d+)\.\s+", t)
            if m:
                num = int(m.group(1))
                if num in (20, 21, 22):
                    ref_paras[num] = par
            continue

        new_text = _fix_in_text_citations(par.text)
        if new_text != par.text:
            _replace_runs_text(par, new_text)

    if set(ref_paras) != {20, 21, 22}:
        raise RuntimeError(f"Expected refs 20-22 in list, found: {sorted(ref_paras)}")

    old_text = {n: ref_paras[n].text for n in (20, 21, 22)}
    new_bodies = {
        20: re.sub(r"^21\.\s+", "20. ", old_text[21], count=1),
        21: re.sub(r"^22\.\s+", "21. ", old_text[22], count=1),
        22: re.sub(r"^20\.\s+", "22. ", old_text[20], count=1),
    }
    for num, body in new_bodies.items():
        _replace_runs_text(ref_paras[num], body)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print("  Body: [21,22]->[20,21], [20]->[22], [22]->[21]")
    print("  References 20=Anselin 1988, 21=LeSage 2009, 22=Anselin LISA 1995")


if __name__ == "__main__":
    main()
