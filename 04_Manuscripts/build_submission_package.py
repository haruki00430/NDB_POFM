#!/usr/bin/env python3
"""Build Health Policy Editorial Manager submission_package/ from project sources."""

from __future__ import annotations

import argparse
import shutil
import subprocess
from pathlib import Path

from docx import Document

PROJECT_DIR = Path(__file__).resolve().parents[1]
MS_DIR = PROJECT_DIR / "04_Manuscripts"
PKG_DIR = MS_DIR / "submission_package"
FIG_DIR = PROJECT_DIR / "03_Analysis" / "results" / "figures"
FIG_ALT = PROJECT_DIR / "results" / "figures"

# Manuscript copy policy:
# - Canonical (手動編集正本): submission_package/manuscript_main.docx
# - Git 管理コピー: 10Manuscript_health_policy_submission_v1_anonymous.docx
MANUSCRIPT_CANONICAL = PKG_DIR / "manuscript_main.docx"
MANUSCRIPT_REPO = MS_DIR / "10Manuscript_health_policy_submission_v1_anonymous.docx"

COPY_MAP: list[tuple[Path, Path]] = [
    (
        MS_DIR / "10b_TitlePage_health_policy.docx",
        PKG_DIR / "title_page.docx",
    ),
]

FIGURE_MAP: list[tuple[Path, Path]] = [
    (FIG_DIR / "choropleth_dental_mgmt_rate.png", PKG_DIR / "Figure3_choropleth_pofm.png"),
    (FIG_DIR / "lisa_cluster_map.png", PKG_DIR / "Figure6_lisa_cluster_map.png"),
    (
        FIG_DIR / "choropleth_cancer_surgery_rate.png",
        PKG_DIR / "AppendixFigureA1_cancer_surgery.png",
    ),
    (
        FIG_DIR / "scatter_surgery_vs_dental.png",
        PKG_DIR / "AppendixFigureA2_scatter_surgery.png",
    ),
    (FIG_DIR / "forest_sensitivity.png", PKG_DIR / "AppendixFigureA4_forest_sensitivity.png"),
    (
        FIG_ALT / "scatter_dentist_vs_dental.png",
        PKG_DIR / "AppendixFigureA5_scatter_dentist.png",
    ),
]


def _run_pandoc(src: Path, dst: Path) -> None:
    """Convert Markdown to DOCX via pandoc."""
    subprocess.run(
        ["pandoc", str(src), "-o", str(dst), "--from", "markdown", "--to", "docx"],
        check=True,
    )


def _extract_highlights(manuscript: Path) -> list[str]:
    """Extract Highlights bullets from manuscript_main.docx."""
    doc = Document(str(manuscript))
    lines: list[str] = []
    in_highlights = False
    for par in doc.paragraphs:
        text = par.text.strip()
        if text == "Highlights":
            in_highlights = True
            continue
        if in_highlights:
            if text.startswith("1. Abstract") or text.startswith("Abstract"):
                break
            if not text:
                continue
            bullet = text.lstrip("•").strip()
            if bullet:
                lines.append(bullet)
    if not lines:
        raise RuntimeError("Highlights section not found in manuscript DOCX")
    return lines


def _write_highlights_txt(lines: list[str], out_txt: Path) -> None:
    """Write plain-text Highlights for reference."""
    out_txt.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _write_highlights_docx(lines: list[str], out_docx: Path) -> None:
    """Write Highlights as a standalone DOCX for EM upload."""
    doc = Document()
    title = doc.add_paragraph()
    title_run = title.add_run("Highlights")
    title_run.bold = True
    for line in lines:
        doc.add_paragraph(f"• {line}")
    doc.save(str(out_docx))


def _resolve_manuscript(force: bool, sync_repo: bool) -> Path:
    """Resolve manuscript path without overwriting the canonical file by default."""
    if sync_repo:
        if not MANUSCRIPT_CANONICAL.exists():
            raise FileNotFoundError(
                f"Cannot sync repo copy; canonical missing: {MANUSCRIPT_CANONICAL}"
            )
        shutil.copy2(MANUSCRIPT_CANONICAL, MANUSCRIPT_REPO)
        print(f"Synced repo copy: {MANUSCRIPT_REPO.name} <- {MANUSCRIPT_CANONICAL.name}")
        return MANUSCRIPT_CANONICAL

    if force:
        if not MANUSCRIPT_REPO.exists():
            raise FileNotFoundError(f"Missing repo manuscript source: {MANUSCRIPT_REPO}")
        shutil.copy2(MANUSCRIPT_REPO, MANUSCRIPT_CANONICAL)
        print(f"Overwritten (--force): {MANUSCRIPT_CANONICAL.name} <- {MANUSCRIPT_REPO.name}")
        return MANUSCRIPT_CANONICAL

    if MANUSCRIPT_CANONICAL.exists():
        print(f"Using canonical (unchanged): {MANUSCRIPT_CANONICAL.name}")
        return MANUSCRIPT_CANONICAL

    if MANUSCRIPT_REPO.exists():
        shutil.copy2(MANUSCRIPT_REPO, MANUSCRIPT_CANONICAL)
        print(f"Initialized: {MANUSCRIPT_CANONICAL.name} <- {MANUSCRIPT_REPO.name}")
        return MANUSCRIPT_CANONICAL

    raise FileNotFoundError(
        "No manuscript found. Place manuscript_main.docx in submission_package/ "
        f"or create {MANUSCRIPT_REPO.name}"
    )


def _write_readme() -> None:
    readme = PKG_DIR / "README_UPLOAD_ORDER.md"
    readme.write_text(
        """# Health Policy — submission_package

ローカル専用の Editorial Manager 投稿ステージングフォルダです。
原稿を更新したら `python build_submission_package.py` で再生成してください。

## 原稿ファイルの正本

| 役割 | パス |
|------|------|
| **正本（手動編集）** | `submission_package/manuscript_main.docx` |
| Git 管理コピー | `04_Manuscripts/10Manuscript_health_policy_submission_v1_anonymous.docx` |

- 通常の `build_submission_package.py` は **manuscript_main.docx を上書きしません**
- `--force` … `10Manuscript_...anonymous.docx` から manuscript_main.docx を上書き
- `--sync-repo` … manuscript_main.docx を 10Manuscript_...anonymous.docx に反映

## ファイル一覧とコピー元

| Package ファイル | コピー元 |
|-----------------|---------|
| `manuscript_main.docx` | **正本（上書きしない）** — Highlights 抽出もここから |
| `title_page.docx` | `04_Manuscripts/10b_TitlePage_health_policy.docx` |
| `cover_letter.docx` | `04_Manuscripts/cover_letter_health_policy.md`（pandoc 変換） |
| `highlights.docx` | manuscript_main.docx の Highlights セクションから生成 |
| `highlights.txt` | 同上（プレーンテキスト参照用） |
| `STROBE_checklist.docx` | `04_Manuscripts/STROBE_checklist_perioperative.md`（pandoc 変換） |
| `Figure3_choropleth_pofm.png` | `03_Analysis/results/figures/choropleth_dental_mgmt_rate.png` |
| `Figure6_lisa_cluster_map.png` | `03_Analysis/results/figures/lisa_cluster_map.png` |
| `AppendixFigureA1_cancer_surgery.png` | `03_Analysis/results/figures/choropleth_cancer_surgery_rate.png` |
| `AppendixFigureA2_scatter_surgery.png` | `03_Analysis/results/figures/scatter_surgery_vs_dental.png` |
| `AppendixFigureA4_forest_sensitivity.png` | `03_Analysis/results/figures/forest_sensitivity.png` |
| `AppendixFigureA5_scatter_dentist.png` | `results/figures/scatter_dentist_vs_dental.png` |

## EM アップロード順（目安）

詳細は `../EM_submission_guide_HEALTH_POLICY.md` を参照。

1. Title Page → `title_page.docx`
2. Manuscript（匿名） → `manuscript_main.docx`
3. Figure → `Figure3_*.png`, `Figure6_*.png`
4. Cover Letter → `cover_letter.docx`
5. Highlights → `highlights.docx`
6. Supplementary / Checklist → `STROBE_checklist.docx`
7. Appendix Figures → `AppendixFigureA*.png`

## 投稿先

- **Health Policy**（Elsevier）
- Editorial Manager: https://www.editorialmanager.com/HEAP/default.aspx
""",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Health Policy submission_package/")
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite manuscript_main.docx from 10Manuscript_...anonymous.docx",
    )
    parser.add_argument(
        "--sync-repo",
        action="store_true",
        help="Copy manuscript_main.docx to 10Manuscript_...anonymous.docx (repo sync)",
    )
    args = parser.parse_args()

    if args.force and args.sync_repo:
        parser.error("Use either --force or --sync-repo, not both.")

    PKG_DIR.mkdir(parents=True, exist_ok=True)

    manuscript = _resolve_manuscript(force=args.force, sync_repo=args.sync_repo)

    for src, dst in COPY_MAP:
        if not src.exists():
            raise FileNotFoundError(f"Missing source: {src}")
        shutil.copy2(src, dst)
        print(f"Copied: {dst.name} <- {src.name}")

    cover_md = MS_DIR / "cover_letter_health_policy.md"
    strobe_md = MS_DIR / "STROBE_checklist_perioperative.md"
    cover_docx = PKG_DIR / "cover_letter.docx"
    strobe_docx = PKG_DIR / "STROBE_checklist.docx"
    _run_pandoc(cover_md, cover_docx)
    print(f"Converted: {cover_docx.name}")
    _run_pandoc(strobe_md, strobe_docx)
    print(f"Converted: {strobe_docx.name}")

    highlight_lines = _extract_highlights(manuscript)
    highlights_txt = PKG_DIR / "highlights.txt"
    highlights_docx = PKG_DIR / "highlights.docx"
    highlights_src = MS_DIR / "highlights_health_policy.docx"
    _write_highlights_txt(highlight_lines, highlights_txt)
    _write_highlights_docx(highlight_lines, highlights_docx)
    _write_highlights_docx(highlight_lines, highlights_src)
    print(f"Generated: {highlights_docx.name} ({len(highlight_lines)} bullets)")
    print(f"Generated: {highlights_src.name}")
    print(f"Extracted: {highlights_txt.name}")

    for src, dst in FIGURE_MAP:
        if not src.exists():
            raise FileNotFoundError(f"Missing figure: {src}")
        shutil.copy2(src, dst)
        print(f"Copied: {dst.name}")

    _write_readme()
    print(f"\nDone: {PKG_DIR}")


if __name__ == "__main__":
    main()
