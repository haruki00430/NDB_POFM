#!/usr/bin/env python3
"""
Health Policy (Elsevier) submission DOCX preparation.
Base: 09Manuscript_perioperative_oral_care_policy_reframed_v9_figures_tables_final.docx
Output: 10Manuscript_health_policy_submission_v1_anonymous.docx

Changes:
  1. Title  → add "in Japan", replace "Study" with "Analysis"
  2. Authors → anonymize (double-blind)
  3. Date   → remove
  4. Highlights → new section before Abstract
  5. Abstract  → add Objectives heading
  6. Research in Context → new section before Introduction
  7. References (n) → [n] in body text
  8. Figure callouts → update for Appendix figures
  9. Tables/Figures section → Main (T1,T2,F3,F6) / Appendix (F1,F2,F4,F5)
"""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = '09Manuscript_perioperative_oral_care_policy_reframed_v9_figures_tables_final.docx'
OUTPUT = '10Manuscript_health_policy_submission_v1_anonymous.docx'

doc = Document(INPUT)
body = doc.element.body
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# ── helpers ──────────────────────────────────────────────────────────────────

def new_para(text='', bold=False):
    p = OxmlElement('w:p')
    if text:
        r = OxmlElement('w:r')
        if bold:
            rpr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rpr.append(b)
            r.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(XML_SPACE, 'preserve')
        r.append(t)
        p.append(r)
    return p

def para_text(elem):
    return ''.join((t.text or '') for t in elem.iter(qn('w:t'))).strip()

def find_in_body(pattern):
    for e in body:
        if e.tag.endswith('}p') and re.search(pattern, para_text(e)):
            return e
    return None

def require_in_body(pattern):
    """Return a body paragraph element or raise if the pattern is missing."""
    elem = find_in_body(pattern)
    if elem is None:
        raise RuntimeError(f"Required section not found: {pattern}")
    return elem

def get_paras():
    return doc.paragraphs

def add_before_end(elems):
    """Insert each element immediately before the final sectPr."""
    sectpr = list(body)[-1]
    for e in elems:
        sectpr.addprevious(e)

# ── 1. TITLE ─────────────────────────────────────────────────────────────────
p = get_paras()[0]
for run in p.runs:
    run.text = run.text.replace(
        'A Nationwide Ecological Study of Supply-Sensitive Implementation in Perioperative Oral Care',
        'A Prefectural Ecological Analysis of Perioperative Oral Care in Japan'
    )
print(f'1. Title → {get_paras()[0].text[:90]}')

# ── 2. ANONYMIZE AUTHORS ──────────────────────────────────────────────────────
for idx in [1, 2]:
    p = get_paras()[idx]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = (
            '[Author details removed for double-anonymized peer review]' if idx == 1 else ''
        )
    else:
        p.add_run('[Author details removed for double-anonymized peer review]' if idx == 1 else '')
print('2. Authors anonymized')

# ── 3. REMOVE DATE ────────────────────────────────────────────────────────────
p = get_paras()[3]
for run in p.runs:
    run.text = ''
print('3. Date removed')

# ── 4. HIGHLIGHTS (inserted before "1. Abstract") ────────────────────────────
abstract_heading = require_in_body(r'^1\.\s+Abstract$')
highlights = [
    (True,  'Highlights'),
    (False, '• POFM delivery varied 2.25-fold across 47 Japanese prefectures'),
    (False, '• Cancer surgical demand was not associated with POFM (0/7 specifications)'),
    (False, '• Dentist density was the strongest predictor of POFM (r = 0.617, p < 0.0001)'),
    (False, '• Significant spatial clustering identified (Moran’s I = 0.210, p = 0.024)'),
    (False, '• Supply-side capacity, not surgical need, drives regional POFM inequity in Japan'),
    (False, ''),
]
# Insert in reverse so final order is correct
prev = abstract_heading
for bold, text in reversed(highlights):
    e = new_para(text, bold=bold)
    prev.addprevious(e)
    prev = e
print('4. Highlights added')

# ── 5. ABSTRACT: add Objectives ───────────────────────────────────────────────
bg_elem = require_in_body(r'^Background:')
obj_text = (
    'Objectives: To characterize prefectural variation in perioperative oral functional management '
    '(POFM) delivery across Japan, examine whether this variation was associated with cancer surgical '
    'demand or with dental workforce supply, and assess spatial clustering of POFM implementation.'
)
bg_elem.addnext(new_para(obj_text))
print('5. Abstract Objectives added')

# ── 6. RESEARCH IN CONTEXT (inserted before "2. Introduction") ───────────────
intro_elem = require_in_body(r'^2\.\s+Introduction$')
ric = [
    (False, ''),
    (True,  'Research in Context'),
    (True,  'What is already known on this topic'),
    (False, (
        'Universal health coverage aims to ensure that evidence-based care reaches patients '
        'according to clinical need. Japan incorporated perioperative oral functional management '
        '(POFM) into its national health insurance fee schedule in 2012, with further reimbursement '
        'expansion in 2020 and 2022. Prior studies have documented postoperative benefits of POFM '
        'at the patient and hospital level, but prefecture-level variation in POFM implementation '
        'had not been characterised.'
    )),
    (True,  'What this study adds'),
    (False, (
        'This nationwide prefectural ecological analysis of Japan’s NDB Open Data (fiscal year 2023) '
        'demonstrates a 2.25-fold variation in POFM claim density across 47 prefectures. Cancer surgical '
        'demand—the primary clinical need indicator—was not significantly associated with POFM '
        'provision in any of seven pre-specified sensitivity specifications. By contrast, dentist density '
        'showed the strongest bivariate association (r = 0.617), and significant positive '
        'spatial autocorrelation was identified (Moran’s I = 0.210, p = 0.024).'
    )),
    (True,  'How this study might affect research, practice or policy'),
    (False, (
        'These findings indicate that reimbursement alone does not guarantee need-aligned implementation '
        'of POFM. Dental workforce capacity, hospital–dental integration, and referral infrastructure '
        'appear to be equally important determinants. For countries introducing perioperative oral care '
        'reimbursement, parallel investment in implementation capacity may be essential to prevent '
        'supply-driven inequity.'
    )),
    (False, ''),
]
prev = intro_elem
for bold, text in reversed(ric):
    e = new_para(text, bold=bold)
    prev.addprevious(e)
    prev = e
print('6. Research in Context added')

# ── 7. REFERENCE FORMAT (n) → [n] in body text ───────────────────────────────
ref_section_elem = find_in_body(r'^11\.\s+References?$')
ref_pattern = re.compile(r'\((\d[\d,\s]*)\)')
past_refs = False
for p in get_paras():
    if p._element is ref_section_elem:
        past_refs = True
    if past_refs:
        break
    for run in p.runs:
        if run.text:
            run.text = ref_pattern.sub(
                lambda m: '[' + m.group(1).strip() + ']', run.text
            )
print('7. Reference format (n) → [n]')

# ── 8. FIGURE CALLOUTS IN BODY TEXT ──────────────────────────────────────────
callout_map = {
    '(See Figure 1)': '(See Appendix Figure A1)',
    '(See Figure 2)': '(See Appendix Figure A2)',
    '(See Figure 4)': '(See Appendix Figure A4)',
    '(See Figure 5)': '(See Appendix Figure A5)',
    '(See Figure 3 and Supplementary Figure S1)': '(See Figure 3 and Appendix Figure A6)',
}
for p in get_paras():
    for run in p.runs:
        if run.text:
            for old, new in callout_map.items():
                run.text = run.text.replace(old, new)
print('8. Figure callouts updated')

# ── 9. RESTRUCTURE TABLES / FIGURES SECTION ──────────────────────────────────
# After steps 4-6 inserted paragraphs, re-read body element list
body_list = list(body)
sectpr = body_list[-1]

def find_idx(body_list, pattern):
    for i, e in enumerate(body_list):
        if e.tag.endswith('}p') and re.search(pattern, para_text(e)):
            return i, e
    return None, None

_, e_13     = find_idx(body_list, r'^13\. Figures$')
idx_13, _   = find_idx(body_list, r'^13\. Figures$')
idx_f1, _   = find_idx(body_list, r'^Figure 1:')
idx_f2, _   = find_idx(body_list, r'^Figure 2:')
idx_f3, _   = find_idx(body_list, r'^Figure 3:')
idx_f4, _   = find_idx(body_list, r'^Figure 4:')
idx_f5, _   = find_idx(body_list, r'^Figure 5:')
idx_f6, _   = find_idx(body_list, r'^Figure 6:')
idx_sf, _   = find_idx(body_list, r'^Supplementary Figures$')
idx_stbl, _ = find_idx(body_list, r'^Supplementary Table S1\.')
idx_sf1, _  = find_idx(body_list, r'^Supplementary Figure S1:')

print(f'  F1={idx_f1} F2={idx_f2} F3={idx_f3} F4={idx_f4} '
      f'F5={idx_f5} F6={idx_f6} SF={idx_sf} ST={idx_stbl} SF1={idx_sf1}')

# Collect blocks as sublists of body_list
fig1_block  = body_list[idx_f1  : idx_f2]
fig2_block  = body_list[idx_f2  : idx_f3]
fig3_block  = body_list[idx_f3  : idx_f4]
fig4_block  = body_list[idx_f4  : idx_f5]
fig5_block  = body_list[idx_f5  : idx_f6]
fig6_block  = body_list[idx_f6  : idx_sf]
sfigs_hdr   = body_list[idx_sf  : idx_sf1]
sfig1_block = body_list[idx_sf1 : -1]           # exclude sectPr
stbl_block  = body_list[idx_stbl: idx_13]        # Supp Table S1 → just before "13. Figures"

# Remove everything from Supp Table S1 onward (keep sectPr)
for e in body_list[idx_stbl : -1]:
    body.remove(e)

# Rename captions in the collected blocks
def rename_caption(block, old_prefix, new_prefix):
    for e in block:
        if e.tag.endswith('}p') and para_text(e).startswith(old_prefix):
            for t in e.iter(qn('w:t')):
                if t.text:
                    t.text = t.text.replace(old_prefix, new_prefix, 1)
            break

rename_caption(stbl_block,  'Supplementary Table S1',   'Appendix Table A1')
rename_caption(fig1_block,  'Figure 1:',                'Appendix Figure A1:')
rename_caption(fig2_block,  'Figure 2:',                'Appendix Figure A2:')
rename_caption(fig4_block,  'Figure 4:',                'Appendix Figure A4:')
rename_caption(fig5_block,  'Figure 5:',                'Appendix Figure A5:')

# Rebuild: Main Figures → F3, F6
add_before_end([new_para('13. Main Figures', bold=True)])
add_before_end(fig3_block)
add_before_end(fig6_block)

# Appendix: Supp Table A1, then Appendix Figures A1,A2,A4,A5, then Supp Fig S1
add_before_end([new_para(''), new_para('Appendix', bold=True)])
add_before_end(stbl_block)
add_before_end([new_para('Appendix Figures', bold=True)])
add_before_end(fig1_block)
add_before_end(fig2_block)
add_before_end(fig4_block)
add_before_end(fig5_block)
add_before_end(sfigs_hdr)
add_before_end(sfig1_block)

print('9. Tables/Figures restructured')

# ── 10. ANONYMIZE AUTHOR CONTRIBUTIONS (end matter) ───────────────────────────
AUTHOR_1 = (
    '[Author 1]: Conceptualization, Data curation, Formal analysis, Investigation, '
    'Methodology, Software, Visualization, Writing – original draft, '
    'Writing – review and editing.'
)
AUTHOR_2 = (
    '[Author 2]: Conceptualization, Supervision, Writing – review and editing.'
)
for p in get_paras():
    t = para_text(p._element)
    if t.startswith('Haruki Saito:'):
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = AUTHOR_1
        else:
            p.add_run(AUTHOR_1)
    elif t.startswith('Tetsuya Ohira:'):
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = AUTHOR_2
        else:
            p.add_run(AUTHOR_2)
print('10. Author Contributions anonymized')

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f'\n✓ Saved: {OUTPUT}')
